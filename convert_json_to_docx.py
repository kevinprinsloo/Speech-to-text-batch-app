import json
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt, RGBColor

# Load JSON data from a file
def load_json(file_path):
    with open(file_path, 'r', encoding='utf-8') as file:
        return json.load(file)

# Define colors for speakers
SPEAKER_COLORS = {
    "speaker_1": RGBColor(0, 0, 255),  # Blue
    "speaker_2": RGBColor(255, 0, 0),  # Red
    # Add more speakers and colors if needed
}

# Create a Word document from JSON data
def create_word_doc(conversation, output_file):
    doc = Document()

    # Add a title
    title = doc.add_heading('Speaker Conversation', level=1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Add the conversation
    for entry in conversation:
        speaker = entry['speaker']
        text = entry['text']
        timestamp = entry['timestamp']

        # Add speaker and text in the same paragraph
        speaker_paragraph = doc.add_paragraph()
        speaker_run = speaker_paragraph.add_run(f"{speaker}: ")
        speaker_run.bold = True
        speaker_run.font.size = Pt(12)
        if speaker in SPEAKER_COLORS:
            speaker_run.font.color.rgb = SPEAKER_COLORS[speaker]

        text_run = speaker_paragraph.add_run(text)
        text_run.font.size = Pt(12)

        # Add timestamp in the same paragraph, but formatted smaller and grey
        timestamp_run = speaker_paragraph.add_run(f" [{timestamp}]")
        timestamp_run.italic = True
        timestamp_run.font.size = Pt(8)
        timestamp_run.font.color.rgb = RGBColor(128, 128, 128)  # Grey
        speaker_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    # Save the document
    doc.save(output_file)

# Main function
def main():
    input_file = "C:/Users/KevinPrinsloo/OneDrive/aa_SKY/Transcription_Audio/output/5751d43c-5c47-445b-b30c-6c829d5c7bc3_speaker_conversation.json"
    output_file = "C:/Users/KevinPrinsloo/OneDrive/aa_SKY/Transcription_Audio/output/Speaker_Conversation.docx"

    conversation_data = load_json(input_file)
    conversation = conversation_data['conversation']
    create_word_doc(conversation, output_file)
    print(f"Word document created successfully at {output_file}")

if __name__ == "__main__":
    main()
